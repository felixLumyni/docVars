'''
# docVars v1.7 by Lumyni
# Requires https://www.python.org/
# Messes w/ files, only edit this if you know what you're doing!
'''

import os, importlib, sys, locale, ctypes, re
from datetime import datetime
from tkinter import messagebox
from tkinter import filedialog

def import_required_modules(modules):
    for (module,link,targetversion,extramodules) in modules:
        parameters = ""
        try:
            moduleobj = __import__(module)
            globals()[module] = moduleobj
            version = moduleobj.__version__.replace('.','')
            if version < targetversion:
                parameters = "--upgrade" 
                print(f"An inferior version of the module '{module}' was found ({version} vs {targetversion}).")
                consent = input("Would you like to continue trying to run the app anyways? NOT RECOMMENDED (Y/n) ")
                if not(consent.upper() == "Y"):
                    print("Operation denied by user.")
                    raise Exception(f"Inferior version ({version} vs {targetversion}).")
                else:
                    print("Oh well. Trying to run the app anyways...")
            elif targetversion != '0' and version > targetversion:
                print(f"WARNING: Current version of '{module}' ({version}) is higher than the one used in this script ({targetversion}).")
        except Exception as reason:
            print(f"Couldn't find the required module '{module}'{link} \nReason: {reason}")
            consent = input(f"Would you like to automatically install it now with 'pip install {module} {extramodules} {parameters}'? (Y/n) ")
            pendingExit = False
            if consent.upper() == "Y":
                print("Operation accepted by user.")
                os.system(f'pip install {module} {extramodules} {parameters}')
                try:
                    moduleobj = __import__(module)
                    globals()[module] = moduleobj
                except:
                    print("\nCouldn't automatically install module, it is likely that this script cannot access pip.")
                    pendingExit = True
            else:
                print("Operation denied by user.")
                pendingExit = True
            if pendingExit == True:
                print(f"Please install {module} before reopening this app.")
                input("(Press enter to quit)\n")
                quit()

def import_path(path):
    module_name = os.path.basename(path).replace('-', '_')
    spec = importlib.util.spec_from_loader(
        module_name,
        importlib.machinery.SourceFileLoader(module_name, path)
    )
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    sys.modules[module_name] = module
    return module

def docx_replace_regex(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def run(onlysave=False, source='', autosave=None, appearancemode=0, language='en', regexes=[], inputs=[]):
    variables = vars()
    if onlysave:
        settings_save(variables, onlysave)
        return
    else:
        now = datetime.now()
        dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
        print(f"\n[{dt_string}] Initiating procedure for {os.path.basename(source)}...")
    try:
        count = 0
        doc = docx.Document(source)
        for r in regexes:
            regex = re.compile(r""+r)
            input = r""+inputs[count]
            docx_replace_regex(doc, regex, input)
            print(f"Replaced all occurrences of {r} with {inputs[count]}")
            count += 1
        extension = os.path.splitext(os.path.basename(source))[1]
        result = source.replace(extension,"_pronto"+extension)
        print(f"Saved result document as {os.path.basename(result)}")
        doc.save(result)
    except Exception as e:
        msg = "Something went wrong.\nRefer to the console for details."
        print(f"ERROR: {e}")
        messagebox.showerror(title='Error (Operation cancelled)', message=msg)
        return
    if autosave: settings_save(variables, onlysave)

def settings_save(variables, settingsfile=False):
    if type(settingsfile) is bool:
        settingsfile = preset
        if preset == "(Default)":
            settingsfile = os.path.basename(__file__).replace(".py","_settings.txt")
    #TODO: don't delete vars that won't be replaced?
    #don't save these
    try: del variables['variables']
    except: pass
    try: del variables['onlysave']
    except: pass
    if variables['appearancemode'] == 0:
        try: del variables['appearancemode']
        except: pass
    #iteration time... is this even readable?
    try:
        settings = f""
        for v in list(variables.items()):
            settings += f"{v[0]} = {repr(v[1])}\n" #variable_name = variable, then go to the next line
        with open(settingsfile, "w") as outfile:
            outfile.write(settings)
            print(f"Saved settings to {os.path.basename(settingsfile)}")
    except Exception as e:
        print(f"WARNING: Couldn't save settings file. Reason: {e}")
        messagebox.showwarning(title='Warning', message="Something went wrong when trying to save the settings file.\nCheck the console for details.")

class UX:
    def __init__(self, root, warn):
        lists = ["@PARAM","?"]
        def renderlists():
            listcount = 0
            elements = self.frame_1.winfo_children()
            for list in lists:
                listcount += 1
                elemcount = 0
                for i in enumerate(elements):
                    if 'ctkentry' in str(i[1]):
                        if elements[i[0]].cget('placeholder_text') and (list in elements[i[0]].cget('placeholder_text')):
                            #magic numbers from loadPage()
                            elemcount += 1
                            elements[i[0]].place(x=57+(130*listcount), y=10+29*(3+elemcount))  
                            self.add.place(y=10+29*(4+elemcount))   
                            self.bar1.configure(height=29*(2+elemcount))
                    elif 'ctkbutton' in str(i[1]) and elements[i[0]].cget('text') == "x":
                        elements[i[0]].place(x=20+(142*(listcount-1)), y=15+29*(3+elemcount))

        def loadPage(mode=None):
            if mode: self.box1.set(mode)
            goToHell()

            OFFSET = 10
            SPACING = 29
            L = 5
            C = 155
            R = 305
            CL = (C+L)/2
            CR = (C+R)/2
            
            self.swi1.place(x=R+C-48, y=5)

            if mode == None: #there's only this one page atm
                self.t1.configure(state=customtkinter.NORMAL)
                self.inf1.place(x=SPACING*3.5, y=00)
                self.b2.place(x=CL, y=SPACING*1)
                self.b1.place(x=CL, y=SPACING*2)
                self.lbl1.place(x=L, y=OFFSET+SPACING*3)
                self.bar1.place(x=C-2, y=OFFSET+SPACING*3)
                self.lbl2.place(x=C+SPACING*2.2, y=OFFSET+SPACING*3)
                self.lbl3.place(x=R+SPACING*2, y=OFFSET+SPACING*3)
                self.add.place(x=C+CR/2+35, y=OFFSET+SPACING*4)
                self.t1.place(x=L, y=OFFSET+SPACING*4)
                self.cfg1.place(x=CR, y=4+SPACING*1)
                renderlists()
                '''offscreen'''
                self.slbl1.place(x=R+C+5, y=SPACING*0)
                self.sb1.place(x=R+C+5,y=SPACING*1)
                self.sb2.place(x=R+C+5,y=SPACING*2)
            else:
                messagebox.showerror(title="FATAL ERROR", message="wtf")
                quit()
        
        def goToHell(): #i hate this
            #teleport labels, entries, buttons and checkboxes to hell
            #i wish i could just make them invisible, oh well
            elements = self.frame_1.winfo_children()
            for i in enumerate(elements):
                for type in ['ctklabel', 'ctkentry', 'ctkbutton', 'ctkcheckbox']:
                    if type in str(i[1]):
                        if 'ctkentry' in str(i[1]) and elements[i[0]].cget('placeholder_text'):
                            continue
                        elements[i[0]].place(y=-666)
                        if 'ctkentry' in str(i[1]): #also disable entries, because TAB...
                            elements[i[0]].configure(state=customtkinter.DISABLED)

        def switchDark(switchswitch=False):
            theme = customtkinter.get_appearance_mode()
            customtkinter.set_appearance_mode("Dark") if theme == "Light" else customtkinter.set_appearance_mode("Light")
            #run(True, *list(getEVERYTHING()))
            if switchswitch: self.swi1.select() if not self.swi1.get() else self.swi1.deselect()

        def resetEntries():
            elements = self.frame_1.winfo_children()
            #delete
            for i in enumerate(elements):
                if 'ctkentry' in str(i[1]):
                    elements[i[0]].delete(0,customtkinter.END)
            #load 
            try: self.t1.insert(0,settings.source)
            except: pass
            try:
                for count in range(len(settings.regexes)):
                    addLists(fill=[settings.regexes[count], settings.inputs[count]])
            except: pass
            try: self.cfg1.select() if settings.autosave else self.cfg1.deselect()
            except: self.cfg1.select()
            try: switchDark(True) if settings.appearancemode else "pass" #why must this be a string?
            except: pass
            try: self.lang.insert(0,settings.language) if settings.language else self.lang.insert(0,"auto")
            except: self.lang.insert(0,"auto")

        #IMPORTANT: this gets the values based in the order of creation of the entries; mind the run() function
        def getEVERYTHING():
            values = []
            elements = self.frame_1.winfo_children()
            #individual values
            for i in enumerate(elements):
                for type in ['ctkentry', 'ctkcheckbox', 'ctkswitch']:
                    if type in str(i[1]):
                        placeholder_text = None
                        try: placeholder_text = elements[i[0]].cget('placeholder_text')
                        except: pass
                        if placeholder_text == None: values.append(elements[i[0]].get())
            #list values
            for list in lists:
                subvalues = []
                for i in enumerate(elements):
                    if 'ctkentry' in str(i[1]):
                        if elements[i[0]].cget('placeholder_text') == list:
                            subvalues.append(elements[i[0]].get())
                values.append(subvalues)
            return values

        def savepreset():
            path = filedialog.asksaveasfilename(defaultextension=".lumy")
            if not path == '':
                run(path, *list(getEVERYTHING()))
                global preset
                preset = path
                self.slbl1.configure(text = f'Current preset: {os.path.basename(preset)}')

        def loadpreset():
            path = filedialog.askopenfilename()
            if not path == '':
                print("Attempting to change settings file...\n")
                root.destroy()
                for after_id in root.tk.eval('after info').split():
                    root.after_cancel(after_id)
                main(path)

        def smartPath(entry):
            try: path = filedialog.askopenfilename(initialdir=entry.get(),filetypes=[("Word file", ".docx")])
            except: path = filedialog.askopenfilename(filetypes=[("Word file", ".docx")])
            if not path == '':
                entry.delete(0,customtkinter.END)
                entry.insert(0,path)

        def deleteList(elements, button):
            button.destroy()
            for e in elements:
                e.destroy()
            renderlists()

        def addLists(args=lists, fill=None):
            count = 0
            elements = []
            try:
                #entry boxes
                for arg in args:
                    lol = customtkinter.CTkEntry(self.frame_1, placeholder_text=arg, placeholder_text_color="#FF0000", width=130)
                    elements.append(lol)
                    try:
                        lol.insert(0, fill[count])
                        count += 1
                    except:
                        count += 1
                        continue
                #detele button
                lmao = customtkinter.CTkButton(self.frame_1, text="x", font=("Rockwell Extra Bold", 12), command= lambda: deleteList(elements, lmao), width=12, height=12)
                #refresh
                renderlists()
            except Exception as e:
                msg = "Something went wrong whilst trying to insert lists into the interface.\nCheck the console for details."
                print("FATAL ERROR:",e)
                messagebox.showerror("FATAL ERROR",msg)
                warn['type'] = 2

        self.frame_0=customtkinter.CTkFrame(master=root)
        self.frame_0.pack(pady=10, padx=10, expand=True, fill="both")
        self.canvas_1=customtkinter.CTkCanvas(self.frame_0,width=999,height=999)
        self.canvas_1.place(x=-5,y=-5)
        self.scroll1=customtkinter.CTkScrollbar(self.frame_0, orientation=customtkinter.VERTICAL, command=self.canvas_1.yview)
        self.scroll1.place(x=444,y=0, height=99*4.5)
        self.canvas_1.configure(yscrollcommand=self.scroll1.set)
        self.canvas_1.bind('<Configure>', lambda e: self.canvas_1.configure(scrollregion = self.canvas_1.bbox("all")))
        def _on_mouse_wheel(event): self.canvas_1.yview_scroll(-1 * int((event.delta / 120)), "units")
        self.canvas_1.bind_all("<MouseWheel>", _on_mouse_wheel)
        self.frame_1=customtkinter.CTkFrame(self.canvas_1,width=999,height=9999)
        self.frame_1.place(x=0,y=0)
        self.canvas_1.create_window((0,0), window=self.frame_1, anchor="nw")
        #
        self.inf1=customtkinter.CTkLabel(self.frame_1, text='Make sure the paths do not require admin', text_color="#696969")
        self.lbl1=customtkinter.CTkButton(self.frame_1, text='Document template', fg_color="#424242", hover_color="#696969", command= lambda: smartPath(self.t1))
        self.lbl2=customtkinter.CTkLabel(self.frame_1, text='Replace...')
        self.lbl3=customtkinter.CTkLabel(self.frame_1, text='with...')
        self.add=customtkinter.CTkButton(self.frame_1, text='+', font=("Rockwell Extra Bold", 12), width=12, height=12, command= addLists) 
        self.t1=customtkinter.CTkEntry(self.frame_1)
        self.b1=customtkinter.CTkButton(self.frame_1, text='Save settings', command= lambda: run(True, *list(getEVERYTHING())))
        self.b2=customtkinter.CTkButton(self.frame_1, text='Generate document', command= lambda: run(False, *list(getEVERYTHING())))
        self.cfg1=customtkinter.CTkCheckBox(master=self.frame_1, text="Save settings after generation")
        self.slbl1=customtkinter.CTkLabel(self.frame_1, text=f'Current preset: {os.path.basename(preset)}')
        self.sb1=customtkinter.CTkButton(self.frame_1, text='Save preset', command= lambda: savepreset())
        self.sb2=customtkinter.CTkButton(self.frame_1, text='Load preset', command= lambda: loadpreset())
        self.swi1=customtkinter.CTkSwitch(self.frame_1, text="", command=switchDark, width=12)
        self.bar1=customtkinter.CTkButton(self.frame_1, text="", width=2, height=100, fg_color="#424242", hover_color="#424242")
        self.lang=customtkinter.CTkEntry(self.frame_1)

        def translate(lang: str):
            wasAuto = False
            if lang == "auto":
                wasAuto = True
                windll = ctypes.windll.kernel32
                windll.GetUserDefaultUILanguage()
                lang = locale.windows_locale[ windll.GetUserDefaultUILanguage() ]
            lang = lang.upper()
            lang = lang.replace('_','')
            lang = lang.replace('-','')
            lang = lang.replace(' ','')
            if lang == "PTBR" or lang == "PT" or lang == "BR":
                self.inf1.configure(text="N\u00E3o use documentos que precisam de admin")
                self.lbl1.configure(text="Documento original")
                self.lbl2.configure(text="Substituir...")
                self.lbl3.configure(text="por...")
                self.b1.configure(text="Salvar configura\u00E7\u00F5es")
                self.b2.configure(text="Gerar documento")
                self.cfg1.configure(text="Salvar ap\u00F3s gera\u00E7\u00E3o")
                self.slbl1.configure(text="Preset atual: "+os.path.basename(preset))
                self.sb1.configure(text="Salvar preset")
                self.sb2.configure(text="Carregar preset")
                return
            #none of the above returned
            if not(wasAuto or lang == "EN"):
                messagebox.showwarning(title='Warning', message='The selected language has not been implemented.')

        #restore entries based on user's settings and load the first accessible page
        resetEntries()
        loadPage()
        translate(self.lang.get())

        try:
            if os.getuid() == 0:
                self.inf1.configure(text="")
        except:
            pass

        if warn:
            if warn['type'] == 1:
                MSG1 = "The settings file was found but could not be loaded."
                MSG2 = "To avoid conflicts, the file in question was renamed to "+warn['newfilename']+"."
                MSG3 = "Check the console for details."
                self.w1 = messagebox.showwarning(title='Warning', message=f"{MSG1}\n{MSG2}\n{MSG3}")
                warn = False
                root.focus_force()
            if warn['type'] == 2:
                root.destroy()

def main(settingsfile=None):
    os.chdir(os.path.dirname(os.path.realpath(__file__))) #paranoia
    global preset
    if settingsfile == None: 
        settingsfile = os.path.basename(__file__).replace(".py","_settings.txt")
        preset = "(Default)"
    else:
        preset = settingsfile

    #Try to load settings file, if there is one, and is valid
    global settings
    warn = {}
    if os.path.exists(settingsfile):
        try:
            settings = import_path(settingsfile)
            print("INFO: Found settings file.\n")
        except Exception as e:
            print(f"WARNING: Couldn't load settings file. Reason: {e}")
            try:
                os.rename(settingsfile, settingsfile.replace(".txt","_INVALID.txt"))
                print("Added '_INVALID' to the end of filename.")
            except:
                os.unlink(settingsfile.replace(".txt","_INVALID.txt"))
                print("Deleted previous '_INVALID' file.")
                os.rename(settingsfile, settingsfile.replace(".txt","_INVALID.txt"))
                print("Added '_INVALID' to the end of filename.")
            warn['type'] = 1 #Settings file could not be loaded.
            warn['newfilename'] = os.path.basename(settingsfile).replace('.txt','_INVALID.txt')
    else: print("INFO: Settings file not found, running anyway.\n")

    #Create the app's window
    root = customtkinter.CTk()
    GUI = UX(root, warn)
    if warn and warn['type'] == 2: quit()
    root.title("docVars")
    root.geometry('475x475')
    '''
    #root.iconbitmap("myIcon.ico")
    #root.resizable(False, False)
    #root.overrideredirect(1)
    #root.attributes("-alpha",1)
    #root.attributes("-topmost", 1)
    '''
    root.mainloop()

if __name__ == "__main__":
    required_modules = [
        #(NAME, LINK, TARGET_VERSION, EXTRAMODULES),
        ('customtkinter', ': https://github.com/TomSchimansky/CustomTkinter/tags', '503', ''),
        ('docx', '', '0', 'python-docx'),
    ]
    import_required_modules(required_modules)
    sys.dont_write_bytecode = True
    customtkinter.set_appearance_mode("dark") # Modes: "System" (standard), "Dark", "Light"
    customtkinter.set_default_color_theme("dark-blue") # Themes: "blue" (standard), "green", "dark-blue"
    main()
else: #so pycharm shuts up
    import docx 
    import customtkinter